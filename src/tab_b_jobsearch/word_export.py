"""
简历导出模块（docx / md / pdf）。

输入是定制流程产出的干净 Markdown（resume_document.render_markdown 的格式）：
    # 姓名
    联系方式一行
    ## 板块标题
    ### 条目标题
    时间行
    - bullet
    **粗体** 行内标记 / ![alt](data:...) 图片

排版目标：跟用户上传的原始简历风格一致——姓名居中加大加粗、联系方式一行、
各板块标题加粗 + 分隔线、bullet 统一符号缩进——而不是一坨纯文本。
"""

import base64
import io
import os
import re

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem, HRFlowable, Image,
)

EXPORT_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "data", "exports")

_BOLD_INLINE = re.compile(r"\*\*(.+?)\*\*")
_CJK_CHAR = re.compile(r"[一-鿿]")
_IMAGE_LINE = re.compile(r"^!\[[^\]]*\]\(([^)]*)\)$")
_DATA_URI = re.compile(r"^data:(image/[a-z.+-]+);base64,(.+)$", re.IGNORECASE)


# ---------------------------------------------------------------------------
# 把 Markdown 拆成结构化"块"，docx / pdf 共用
# ---------------------------------------------------------------------------
def _parse_blocks(content: str) -> list[dict]:
    """
    返回块列表，每块 {"type": ..., ...}：
      title(姓名) / contact / section(板块标题) / entry(条目标题) / meta(时间等)
      / bullet / body / image
    """
    blocks: list[dict] = []
    lines = content.split("\n")
    prev_type = None
    for raw in lines:
        line = raw.strip()
        if not line:
            prev_type = None
            continue

        img = _IMAGE_LINE.match(line)
        if img:
            blocks.append({"type": "image", "src": img.group(1)})
            prev_type = "image"
            continue

        if line.startswith("# "):
            blocks.append({"type": "title", "text": line[2:].strip()})
            prev_type = "title"
        elif line.startswith("## "):
            blocks.append({"type": "section", "text": line[3:].strip().replace("*", "")})
            prev_type = "section"
        elif line.startswith("### "):
            blocks.append({"type": "entry", "text": line[4:].strip().replace("*", "")})
            prev_type = "entry"
        elif line.startswith("- ") or (line.startswith("* ") and not line.startswith("**")):
            blocks.append({"type": "bullet", "text": line[2:].strip()})
            prev_type = "bullet"
        elif prev_type == "title":
            blocks.append({"type": "contact", "text": line})
            prev_type = "contact"
        elif prev_type == "entry":
            blocks.append({"type": "meta", "text": line})
            prev_type = "meta"
        else:
            blocks.append({"type": "body", "text": line})
            prev_type = "body"
    return blocks


def _decode_image(src: str) -> bytes | None:
    m = _DATA_URI.match(src.strip())
    if m:
        try:
            return base64.b64decode(m.group(2))
        except Exception:
            return None
    if os.path.isfile(src):
        try:
            with open(src, "rb") as f:
                return f.read()
        except OSError:
            return None
    return None


def _md_to_html(text: str) -> str:
    return _BOLD_INLINE.sub(r"<b>\1</b>", text)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------
def _docx_bottom_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "666666")
    borders.append(bottom)
    p_pr.append(borders)


def _docx_runs(paragraph, text: str) -> None:
    for part in re.split(r"(\*\*.*?\*\*)", text):
        if not part:
            continue
        run = paragraph.add_run(part[2:-2] if part.startswith("**") and part.endswith("**") else part)
        if part.startswith("**") and part.endswith("**"):
            run.bold = True


def export_resume_to_docx(content: str, candidate_name: str) -> str:
    os.makedirs(EXPORT_DIR, exist_ok=True)
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.6)
        section.left_margin = section.right_margin = Inches(0.75)

    blocks = _parse_blocks(content)
    if not any(b["type"] == "title" for b in blocks) and candidate_name.strip():
        blocks.insert(0, {"type": "title", "text": candidate_name.strip()})

    for b in blocks:
        t = b["type"]
        if t == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(b["text"])
            r.bold = True
            r.font.size = Pt(20)
        elif t == "contact":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(b["text"])
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            p.paragraph_format.space_after = Pt(6)
        elif t == "section":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(b["text"].upper())
            r.bold = True
            r.font.size = Pt(12)
            _docx_bottom_border(p)
        elif t == "entry":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(0)
            _docx_runs(p, b["text"])
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(11)
        elif t == "meta":
            p = doc.add_paragraph()
            r = p.add_run(b["text"])
            r.italic = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        elif t == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            _docx_runs(p, b["text"])
        elif t == "image":
            data = _decode_image(b["src"])
            if data:
                try:
                    doc.add_picture(io.BytesIO(data), width=Inches(1.4))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass
        else:  # body
            p = doc.add_paragraph()
            _docx_runs(p, b["text"])

    safe_name = re.sub(r"[^\w\-]", "_", candidate_name) or "resume"
    filepath = os.path.join(EXPORT_DIR, f"{safe_name}_Resume_Tailored.docx")
    doc.save(filepath)
    return filepath


# ---------------------------------------------------------------------------
# MD
# ---------------------------------------------------------------------------
def export_resume_to_md(content: str, candidate_name: str) -> str:
    os.makedirs(EXPORT_DIR, exist_ok=True)
    safe_name = re.sub(r"[^\w\-]", "_", candidate_name) or "resume"
    filepath = os.path.join(EXPORT_DIR, f"{safe_name}_Resume_Tailored.md")

    body = content.strip()
    if not re.match(r"^\s*#\s", body) and candidate_name.strip():
        body = f"# {candidate_name.strip()}\n\n{body}"
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(body + "\n")
    return filepath


# ---------------------------------------------------------------------------
# PDF —— reportlab Platypus 自己排版（xhtml2pdf 的中文渲染有缺陷，见 git 历史）
# ---------------------------------------------------------------------------
_CJK_FONT_CANDIDATES = [
    r"C:\Windows\Fonts\simhei.ttf",
    r"C:\Windows\Fonts\msyh.ttc",
    "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/System/Library/Fonts/PingFang.ttc",
]
_cjk_font_registered = False


def _register_cjk_font() -> str:
    global _cjk_font_registered
    if _cjk_font_registered:
        return "CJKFont"
    for path in _CJK_FONT_CANDIDATES:
        if os.path.exists(path):
            pdfmetrics.registerFont(TTFont("CJKFont", path))
            _cjk_font_registered = True
            return "CJKFont"
    return "Helvetica"


def _font_for(text: str, cjk_font: str, bold: bool = False) -> str:
    if _CJK_CHAR.search(text):
        return cjk_font
    return "Helvetica-Bold" if bold else "Helvetica"


def export_resume_to_pdf(content: str, candidate_name: str) -> str:
    os.makedirs(EXPORT_DIR, exist_ok=True)
    safe_name = re.sub(r"[^\w\-]", "_", candidate_name) or "resume"
    filepath = os.path.join(EXPORT_DIR, f"{safe_name}_Resume_Tailored.pdf")
    cjk = _register_cjk_font()

    blocks = _parse_blocks(content)
    if not any(b["type"] == "title" for b in blocks) and candidate_name.strip():
        blocks.insert(0, {"type": "title", "text": candidate_name.strip()})

    story = []
    bullet_buf: list[str] = []

    def _flush_bullets():
        if not bullet_buf:
            return
        items = [
            ListItem(Paragraph(_md_to_html(b), ParagraphStyle(
                "Bullet", fontName=_font_for(b, cjk), fontSize=10, leading=13.5, spaceAfter=2,
            )))
            for b in bullet_buf
        ]
        story.append(ListFlowable(items, bulletType="bullet", start="•", leftIndent=14))
        bullet_buf.clear()

    for b in blocks:
        t = b["type"]
        if t == "bullet":
            bullet_buf.append(b["text"])
            continue
        _flush_bullets()

        if t == "title":
            story.append(Paragraph(b["text"], ParagraphStyle(
                "Name", fontName=_font_for(b["text"], cjk, bold=True),
                fontSize=20, leading=24, alignment=TA_CENTER,
            )))
        elif t == "contact":
            story.append(Paragraph(b["text"], ParagraphStyle(
                "Contact", fontName=_font_for(b["text"], cjk), fontSize=9.5, leading=12,
                alignment=TA_CENTER, textColor=colors.HexColor("#444444"), spaceAfter=6,
            )))
        elif t == "section":
            story.append(Spacer(1, 8))
            story.append(Paragraph(b["text"].upper(), ParagraphStyle(
                "Section", fontName=_font_for(b["text"], cjk, bold=True),
                fontSize=12, leading=15, spaceAfter=3,
            )))
            story.append(HRFlowable(width="100%", thickness=0.75, color=colors.HexColor("#666666"),
                                    spaceBefore=0, spaceAfter=5))
        elif t == "entry":
            story.append(Spacer(1, 5))
            story.append(Paragraph(_md_to_html(b["text"]), ParagraphStyle(
                "Entry", fontName=_font_for(b["text"], cjk, bold=True), fontSize=11, leading=14,
            )))
        elif t == "meta":
            story.append(Paragraph(b["text"], ParagraphStyle(
                "Meta", fontName=_font_for(b["text"], cjk), fontSize=9.5, leading=12,
                textColor=colors.HexColor("#555555"),
            )))
        elif t == "image":
            data = _decode_image(b["src"])
            if data:
                try:
                    img = Image(io.BytesIO(data))
                    ratio = (img.imageHeight / img.imageWidth) if img.imageWidth else 1
                    img.drawWidth = 1.4 * inch
                    img.drawHeight = 1.4 * inch * ratio
                    img.hAlign = "CENTER"
                    story.append(img)
                except Exception:
                    pass
        else:  # body
            story.append(Paragraph(_md_to_html(b["text"]), ParagraphStyle(
                "Body", fontName=_font_for(b["text"], cjk), fontSize=10, leading=13.5,
            )))

    _flush_bullets()

    SimpleDocTemplate(
        filepath, pagesize=LETTER,
        leftMargin=0.75 * inch, rightMargin=0.75 * inch,
        topMargin=0.6 * inch, bottomMargin=0.6 * inch,
    ).build(story)
    return filepath


if __name__ == "__main__":
    sample = """# Jane Q. Candidate

San Francisco, CA | +1 555 0100 | jane@example.com | github.com/janeq

## PROFESSIONAL SUMMARY

- Backend engineer with 5 years building data-intensive services.

## TECHNICAL SKILLS

- **Languages:** Python, Go, SQL

## PROFESSIONAL EXPERIENCE

### Acme Corp | Payments Platform | Senior Engineer
Jan 2022 - Present

- Cut checkout latency by half by redesigning the settlement pipeline.
- Mentored three engineers on distributed systems.

## PERSONAL PROJECT

### Open Source Rate Limiter
2023

- Built a token-bucket limiter used by 400+ repos.

## EDUCATION

**State University** | B.S. Computer Science | 2013 - 2017
"""
    for fn in (export_resume_to_docx, export_resume_to_md, export_resume_to_pdf):
        print(fn.__name__, "->", fn(sample, "Jane Q. Candidate"))
